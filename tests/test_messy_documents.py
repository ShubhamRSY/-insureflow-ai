"""Messy broker-document intake: xlsx/csv/docx/eml/html parsing + US normalization.

These cover the "Very High" document-processing gaps for US underwriters:
spreadsheet, Word, raw email, and HTML submissions must produce real text and
structured fields instead of UTF-8 garbage from a naive binary decode.
"""

from __future__ import annotations

import base64
import io
import zipfile
from email.message import EmailMessage

from insureflow.ingestion.insurance.loader import InsuranceDocumentLoader
from insureflow.ingestion.insurance.value_normalizers import normalize_amount, normalize_date
from insureflow.ingestion.structured_docs import (
    parse_csv_bytes,
    parse_docx_bytes,
    parse_eml_bytes,
    parse_html_bytes,
    parse_workbook_bytes,
)

DOCX_CONTENT_TYPES = b"""<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">
<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>
<Default Extension="xml" ContentType="application/xml"/>
<Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>
</Types>"""


def _make_docx(*, paragraphs: list[str], table: list[list[str]] | None = None) -> bytes:
    """Build a spec-compliant .docx via python-docx (real broker files look like this)."""
    from docx import Document

    doc = Document()
    for p in paragraphs:
        doc.add_paragraph(p)
    if table:
        t = doc.add_table(rows=len(table), cols=len(table[0]))
        for i, row in enumerate(table):
            for j, cell in enumerate(row):
                t.cell(i, j).text = cell
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_docx_minimal(*, paragraphs: list[str], table: list[list[str]] | None = None) -> bytes:
    """Minimal zip+xml docx — exercises the zero-dependency fallback path."""
    body = '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body>'
    for p in paragraphs:
        body += f'<w:p><w:r><w:t xml:space="preserve">{p}</w:t></w:r></w:p>'
    if table:
        body += "<w:tbl>"
        for row in table:
            body += "<w:tr>" + "".join(f'<w:tc><w:p><w:r><w:t xml:space="preserve">{c}</w:t></w:r></w:p></w:tc>' for c in row) + "</w:tr>"
        body += "</w:tbl>"
    body += "</w:body></w:document>"
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", DOCX_CONTENT_TYPES)
        zf.writestr("word/document.xml", body)
    return buf.getvalue()


def _make_xlsx(sheets: dict[str, list[list[object]]]) -> bytes:
    from openpyxl import Workbook

    wb = Workbook()
    wb.remove(wb.active)
    for name, rows in sheets.items():
        ws = wb.create_sheet(name)
        for row in rows:
            ws.append(row)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _make_eml(*, subject: str, body: str, html: str | None = None) -> bytes:
    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"] = "broker@underwriters.test"
    msg["To"] = "uw@carrier.test"
    msg["Date"] = "Tue, 04 Aug 2026 09:30:00 -0500"
    if html:
        msg.set_content(html, subtype="html")
    else:
        msg.set_content(body)
    return bytes(msg)


def _b64(text: str | bytes) -> str:
    return base64.b64encode(text if isinstance(text, bytes) else text.encode("utf-8")).decode("ascii")


# --------------------------------------------------------------------------- #
# Structured parsers
# --------------------------------------------------------------------------- #
def test_parse_workbook_bytes_markdown_and_fields() -> None:
    data = _make_xlsx(
        {
            "Locations": [
                ["Location", "Building Value", "Contents", "BI"],
                ["1 Main St, Austin TX", 2500000, 500000, 1000000],
                ["400 Oak Ave, Dallas TX", 1750000, 300000, 800000],
            ],
            "Notes": [["Roof 2023"]],
        }
    )
    text, fields = parse_workbook_bytes(data, "sov.xlsx")
    assert "Locations" in text
    assert "2,500,000" in text or "2500000" in text
    assert "Notes" in text
    assert fields["excel.sheets"][0].value == "Locations,Notes"
    assert "excel.unparsed_sheets" not in fields


def test_parse_workbook_blank_sheet_is_unparsed() -> None:
    data = _make_xlsx({"Data": [["a", "b"], ["1", "2"]], "Empty": []})
    text, fields = parse_workbook_bytes(data, "wb.xlsx")
    assert fields["excel.unparsed_sheets"][0].value == "Empty"
    assert "Empty" not in text


def test_parse_csv_bytes_utf8_sig_and_markdown() -> None:
    raw = 'Location,Building Value\n"1 Main St","$2,500,000"\n'.encode("utf-8-sig")
    text, fields = parse_csv_bytes(raw, "sov.csv")
    assert "Location" in text and "2,500,000" in text
    assert fields == {}


def test_parse_csv_bytes_latin1_fallback() -> None:
    raw = "Nombre,Valor\nCafé,1500\n".encode("latin-1")
    text, _ = parse_csv_bytes(raw, "data.csv")
    assert "Café" in text


def test_parse_docx_paragraphs_and_table() -> None:
    data = _make_docx(
        paragraphs=["Insured: Acme Widgets", "NAICS 332913"],
        table=[["Peril", "Limit"], ["Property", "5,000,000"], ["GL", "2,000,000"]],
    )
    text, fields = parse_docx_bytes(data, "submission.docx")
    assert "Insured: Acme Widgets" in text
    assert "NAICS 332913" in text
    assert "Peril" in text and "5,000,000" in text
    assert fields == {}


def test_parse_docx_rejects_non_docx() -> None:
    import pytest

    with pytest.raises(ValueError):
        parse_docx_bytes(b"not a zip at all", "fake.docx")


def test_parse_docx_zipxml_fallback_on_minimal_package() -> None:
    """A stripped OPC package (no .rels) defeats python-docx but the zip+xml fallback reads it."""
    data = _make_docx_minimal(paragraphs=["Insured: Acme Widgets"])
    text, fields = parse_docx_bytes(data, "minimal.docx")
    assert "Insured: Acme Widgets" in text
    assert fields == {}


def test_workbook_numbers_render_without_float_noise() -> None:
    data = _make_xlsx({"Schedule of Values": [["Location", "Building"], ["1 Main St", 2500000]]})
    text, _ = parse_workbook_bytes(data, "sov.xlsx")
    assert "2500000" in text
    assert "2500000.0" not in text


def test_structured_parsers_route_legacy_xls() -> None:
    from insureflow.ingestion.structured_docs import STRUCTURED_PARSERS

    assert STRUCTURED_PARSERS[".xls"] is parse_workbook_bytes
    assert STRUCTURED_PARSERS[".docx"] is parse_docx_bytes
    assert STRUCTURED_PARSERS[".eml"] is parse_eml_bytes


def test_pdfplumber_table_extraction() -> None:
    pytest = __import__("pytest")
    pytest.importorskip("pdfplumber")
    import tempfile
    from pathlib import Path

    try:
        import pymupdf
    except ImportError:
        pymupdf = pytest.importorskip("fitz")

    from insureflow.ingestion.ocr import OCRProcessor

    doc = pymupdf.open()
    page = doc.new_page(width=200, height=120)
    x0, y0, x1, y1 = 20, 20, 180, 100
    cols = [x0, 90, x1]
    rows = [y0, 45, 70, y1]
    for x in cols:
        page.draw_line(pymupdf.Point(x, y0), pymupdf.Point(x, y1))
    for y in rows:
        page.draw_line(pymupdf.Point(x0, y), pymupdf.Point(x1, y))
    cells = [("Location", "Limit"), ("1 Main St", "$2,500,000"), ("400 Oak Ave", "$1,750,000")]
    for i, row in enumerate(cells):
        for j, value in enumerate(row):
            page.insert_text((cols[j] + 4, rows[i] + 14), value)
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(doc.tobytes())
        tmp_path = tmp.name
    doc.close()
    try:
        tables = OCRProcessor()._extract_pdf_tables(tmp_path)
        assert "Location" in tables
        assert "Limit" in tables
        assert "$2,500,000" in tables
    finally:
        Path(tmp_path).unlink(missing_ok=True)


def test_parse_eml_headers_and_body() -> None:
    data = _make_eml(subject="Acme Submission - GL Renewal", body="Attached please find the SOV and loss run.")
    text, fields = parse_eml_bytes(data, "submission.eml")
    assert "Subject: Acme Submission - GL Renewal" in text
    assert "broker@underwriters.test" in text
    assert "Attached please find the SOV" in text
    assert fields["eml.subject"][0].value == "Acme Submission - GL Renewal"


def test_parse_eml_html_body_preferred_fallback() -> None:
    data = _make_eml(
        subject="HTML Quote",
        body="plain version",
        html="<html><body><h1>Quoted</h1><p>Limit <b>$2,500,000</b></p></body></html>",
    )
    text, _ = parse_eml_bytes(data, "q.eml")
    assert "Limit" in text
    assert "$2,500,000" in text


def test_parse_html_strips_markup() -> None:
    text, fields = parse_html_bytes(b"<html><head><style>x</style></head><body><h1>Submission</h1><p>Roof replaced 2023</p></body></html>")
    assert "Submission" in text
    assert "Roof replaced 2023" in text
    assert "<h1>" not in text
    assert fields == {}


# --------------------------------------------------------------------------- #
# Loader integration (binary dispatch + classification)
# --------------------------------------------------------------------------- #
def test_loader_parses_xlsx_into_sov() -> None:
    data = _make_xlsx({"Schedule of Values": [["Location", "Building", "Contents"], ["1 Main St", 2500000, 500000]]})
    loader = InsuranceDocumentLoader()
    bundle = loader.load_from_documents([{"filename": "sov.xlsx", "content": _b64(data), "encoding": "base64"}])
    assert bundle.status.value == "parsed"
    assert len(bundle.unstructured) == 1
    sub = bundle.unstructured[0]
    assert sub.document_type == "schedule_of_values"
    assert sub.raw_text and "2500000" in sub.raw_text
    assert "excel.sheets" in sub.extracted_fields


def test_loader_parses_csv_as_financial_or_supplemental() -> None:
    data = "revenue,expenses,net_income\n1500000,900000,600000\n".encode("utf-8")
    loader = InsuranceDocumentLoader()
    bundle = loader.load_from_documents([{"filename": "fin.csv", "content": _b64(data), "encoding": "base64"}])
    assert bundle.status.value == "parsed"
    assert bundle.unstructured[0].raw_text and "1500000" in bundle.unstructured[0].raw_text


def test_loader_parses_docx_via_base64() -> None:
    data = _make_docx(paragraphs=["Named Insured: Acme", "General Liability 1,000,000"])
    loader = InsuranceDocumentLoader()
    bundle = loader.load_from_documents([{"filename": "note.docx", "content": _b64(data), "encoding": "base64"}])
    sub = bundle.unstructured[0]
    assert "Named Insured: Acme" in sub.raw_text
    assert any(f.field_name == "ocr_engine" and "structured:" in f.value for f in sub.extracted_fields.get("ocr_engine", []))


def test_loader_parses_eml_body() -> None:
    data = _make_eml(subject="Acme - DC 2026-07-01", body="SOV attached, TIV 5,000,000")
    loader = InsuranceDocumentLoader()
    bundle = loader.load_from_documents([{"filename": "acme.eml", "content": _b64(data), "encoding": "base64"}])
    sub = bundle.unstructured[0]
    assert "Subject: Acme - DC" in sub.raw_text
    assert "SOV attached" in sub.raw_text
    assert any(f.field_name == "eml.subject" for f in sub.extracted_fields.get("eml.subject", []))


def test_loader_paths_reads_binary_files() -> None:
    import tempfile
    from pathlib import Path

    with tempfile.TemporaryDirectory() as tmp:
        path = Path(tmp) / "sov.xlsx"
        path.write_bytes(_make_xlsx({"Locations": [["Location", "Building"], ["9 Elm St", 1200000]]}))
        loader = InsuranceDocumentLoader()
        bundle = loader.load_from_paths([str(path)])
        assert bundle.status.value == "parsed"
        assert "1200000" in bundle.unstructured[0].raw_text


# --------------------------------------------------------------------------- #
# US / INR normalization
# --------------------------------------------------------------------------- #
def test_normalize_amount_usd_variants() -> None:
    assert normalize_amount("US$750K") == "750000"
    assert normalize_amount("USD 1.25M") == "1250000"
    assert normalize_amount("$1,250,000.00") == "1250000"
    assert normalize_amount("1.5 million") == "1500000"
    assert normalize_amount("12bn") == "12000000000"


def test_normalize_amount_inr_variants() -> None:
    assert normalize_amount("₹5,00,000") == "500000"
    assert normalize_amount("Rs. 2,50,000") == "250000"
    assert normalize_amount("INR 10,00,000") == "1000000"


def test_normalize_date_us_formats() -> None:
    assert normalize_date("05/24/2026") == "2026-05-24"
    assert normalize_date("5/24/2026") == "2026-05-24"
    assert normalize_date("05.24.2026") == "2026-05-24"
    assert normalize_date("20260524") == "2026-05-24"
    assert normalize_date("2026/05/24") == "2026-05-24"
    assert normalize_date("2026-05-24 14:30:00") == "2026-05-24"


def test_normalize_us_identifiers_preserved() -> None:
    from insureflow.ingestion.insurance.value_normalizers import normalize_field

    assert normalize_field("ein", "12-3456789") == "12-3456789"
    assert normalize_field("phone", "(512) 555-0142") == "(512) 555-0142"
    assert normalize_field("naics_code", "332913") == "332913"
