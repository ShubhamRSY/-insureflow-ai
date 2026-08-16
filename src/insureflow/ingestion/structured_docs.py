"""Structured document extraction for messy broker formats.

Broker submissions rarely arrive as a single clean PDF: spreadsheets, Word
files, raw .eml captures, and HTML printouts are the norm for US underwriters.
These extractors turn each format into plain text (plus extracted fields) using
only core dependencies (openpyxl, lxml, stdlib), so the rest of the pipeline
(classifier, reconciliation, triage) sees real content instead of the UTF-8
``errors="replace"`` garbage a naive decode of binary uploads produces.

Optional libraries (BeautifulSoup) are used when present and skipped otherwise,
keeping the core dependency footprint unchanged.
"""

from __future__ import annotations

import csv
import io
import re
import zipfile
from email import message_from_bytes
from email.message import Message
from pathlib import Path
from typing import Any, Iterator, Sequence, cast

from insureflow.models.submissions import ExtractedField

try:  # pragma: no cover - exercised when optional dep is installed
    from bs4 import BeautifulSoup
except ImportError:  # pragma: no cover
    BeautifulSoup = None  # type: ignore[assignment,misc]

try:  # pragma: no cover - exercised when optional dep is installed
    from lxml import etree
except ImportError:  # pragma: no cover
    etree = None

try:  # pragma: no cover - exercised when optional dep is installed
    from openpyxl import load_workbook
except ImportError:  # pragma: no cover
    load_workbook = None

try:  # pragma: no cover - exercised when optional dep is installed
    from python_calamine import CalamineWorkbook
except ImportError:  # pragma: no cover
    CalamineWorkbook = None  # type: ignore[assignment,misc]

try:  # pragma: no cover - exercised when optional dep is installed
    import docx as _python_docx
except ImportError:  # pragma: no cover
    _python_docx = None  # type: ignore[assignment]

try:  # pragma: no cover - exercised when optional dep is installed
    import polars as _pl
except ImportError:  # pragma: no cover
    _pl = None  # type: ignore[assignment]

_WML = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

_HTML_TAG_RE = re.compile(r"<[^>]+>")
_WHITESPACE_RE = re.compile(r"[ \t\f\v]+")
_MULTI_BLANK_RE = re.compile(r"\n{3,}")


def _collapse_whitespace(text: str) -> str:
    text = _WHITESPACE_RE.sub(" ", text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _MULTI_BLANK_RE.sub("\n\n", text)
    return text.strip()


def _cell_str(value: Any) -> str:
    """Render a cell value without float noise (2500000.0 -> 2500000)."""
    if value is None:
        return ""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value).strip()


def _field(name: str, value: str, context: str = "") -> ExtractedField:
    return ExtractedField(field_name=name, value=str(value), confidence=1.0, context=context)


def _markdown_table(headers: Sequence[str], rows: Sequence[Sequence[Any]]) -> str:
    if not rows:
        return ""
    header = " | ".join(_cell_str(h) for h in headers)
    lines = ["| " + header + " |", "|" + " --- |" * len(headers)]
    for row in rows:
        cells = list(row) + [""] * (len(headers) - len(row))
        lines.append("| " + " | ".join(_cell_str(c).replace("\n", " ") for c in cells[: len(headers)]) + " |")
    return "\n".join(lines)


# --------------------------------------------------------------------------- #
# Spreadsheets (xlsx / xls) and CSV
# --------------------------------------------------------------------------- #
def parse_workbook_bytes(data: bytes, filename: str = "workbook.xlsx") -> tuple[str, dict[str, list[ExtractedField]]]:
    """Extract every sheet of an Excel workbook as markdown + parsed fields.

    Polars (columnar fast path) is preferred when installed, then
    python-calamine handles both legacy ``.xls`` and ``.xlsx``; openpyxl is the
    final fallback. Returns ``(markdown, fields)`` where fields carry keys
    ``excel.sheets`` and ``excel.unparsed_sheets`` so downstream SOV and
    financial-statement parsers can target sheet content.
    """
    polars_result = _parse_workbook_polars(data, filename)
    if polars_result is not None:
        return polars_result
    if CalamineWorkbook is not None:
        return _parse_workbook_calamine(data, filename)
    if load_workbook is None:
        raise ImportError("python-calamine or openpyxl is required for workbook parsing")
    return _parse_workbook_openpyxl(data, filename)


def _parse_workbook_polars(data: bytes, filename: str) -> tuple[str, dict[str, list[ExtractedField]]] | None:
    """Polars fast path — multi-sheet read straight into columnar frames."""
    if _pl is None:
        return None
    try:
        frames = cast(dict[str, Any], _pl.read_excel(io.BytesIO(data), sheet_name=None, engine="calamine"))
    except Exception:
        return None
    sheets: dict[str, list[list[Any]]] = {}
    for title, frame in frames.items():
        rows = [[str(c) if c is not None else "" for c in row] for row in frame.iter_rows()]
        if not rows:
            continue
        header = [str(h) or f"column_{i + 1}" for i, h in enumerate(frame.columns)]
        rows.insert(0, header)
        sheets[title] = rows
    if not sheets:
        return None
    return _render_workbook(sheets, filename)


def _parse_workbook_openpyxl(data: bytes, filename: str) -> tuple[str, dict[str, list[ExtractedField]]]:
    wb = load_workbook(io.BytesIO(data), data_only=True)
    sheets: dict[str, list[list[Any]]] = {}
    for sheet in wb.worksheets:
        sheets[sheet.title] = [list(row) for row in sheet.iter_rows(values_only=True) if any(cell is not None for cell in row)]
    wb.close()
    return _render_workbook(sheets, filename)


def _parse_workbook_calamine(data: bytes, filename: str) -> tuple[str, dict[str, list[ExtractedField]]]:
    wb = CalamineWorkbook.from_filelike(io.BytesIO(data))
    try:
        sheets = {name: wb.get_sheet_by_name(name).to_python() for name in wb.sheet_names}
    finally:
        wb.close()
    return _render_workbook(sheets, filename)


def _render_workbook(sheets: dict[str, list[list[Any]]], filename: str) -> tuple[str, dict[str, list[ExtractedField]]]:
    blocks: list[str] = []
    unparsed: list[str] = []
    for title, rows in sheets.items():
        rows = [row for row in rows if any(cell is not None and str(cell).strip() for cell in row)]
        if not rows:
            unparsed.append(title)
            continue
        blocks.append(f"### Sheet: {title}")
        headers = [_cell_str(h) or f"column_{i + 1}" for i, h in enumerate(rows[0])]
        blocks.append(_markdown_table(headers, rows[1:]))
    fields: dict[str, list[ExtractedField]] = {
        "excel.sheets": [_field("excel.sheets", ",".join(sheets), filename)],
    }
    if unparsed:
        fields["excel.unparsed_sheets"] = [_field("excel.unparsed_sheets", ",".join(unparsed), filename)]
    return "\n\n".join(blocks).strip(), fields


def parse_csv_bytes(data: bytes, filename: str = "file.csv") -> tuple[str, dict[str, list[ExtractedField]]]:
    """Parse CSV bytes into markdown, trying utf-8 then latin-1."""
    for encoding in ("utf-8-sig", "latin-1"):
        try:
            text = data.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:  # pragma: no cover - latin-1 never fails
        text = data.decode("utf-8", errors="replace")
    try:
        rows = [row for row in csv.reader(io.StringIO(text)) if any(cell.strip() for cell in row)]
    except csv.Error:
        return text.strip(), {}
    if not rows:
        return text.strip(), {}
    headers = [h.strip() for h in rows[0]]
    width = max(len(r) for r in rows)
    while len(headers) < width:
        headers.append(f"column_{len(headers) + 1}")
    return _markdown_table(headers, rows[1:]), {}


# --------------------------------------------------------------------------- #
# Word documents (docx) — zero-dependency via zipfile + XML
# --------------------------------------------------------------------------- #
def _docx_text_element(el: Any) -> Iterator[str]:
    tag = el.tag
    if tag == _WML + "t":
        yield el.text or ""
    elif tag in {_WML + "p", _WML + "r", _WML + "ins", _WML + "del"}:
        for child in el:
            yield from _docx_text_element(child)
    elif tag == _WML + "tab":
        yield "\t"


def _docx_paragraph_text(p: Any) -> str:
    return "".join(_docx_text_element(p)).strip()


def _docx_table_markdown(tbl: Any) -> str:
    rows: list[list[str]] = []
    for tr in tbl.iter(_WML + "tr"):
        cells: list[str] = []
        for tc in tr.findall(_WML + "tc"):
            parts = [_docx_paragraph_text(child) for child in tc if child.tag == _WML + "p"]
            cells.append(" ".join(parts))
        rows.append(cells)
    return _markdown_table(rows[0] if rows else [], rows[1:])


def parse_docx_bytes(data: bytes, filename: str = "file.docx") -> tuple[str, dict[str, list[ExtractedField]]]:
    """Extract paragraphs and tables from a .docx (Office Open XML) file.

    python-docx handles nested tables, text boxes, and headers more robustly;
    the zipfile+lxml extractor is the fallback when python-docx is unavailable.
    """
    if _python_docx is not None:
        try:
            return _parse_docx_python_docx(data, filename)
        except Exception:
            pass  # malformed/incomplete OPC package — fall back to zip+xml
    if etree is None:
        raise ImportError("python-docx or lxml is required for .docx parsing")
    return _parse_docx_zipxml(data, filename)


def _parse_docx_python_docx(data: bytes, filename: str) -> tuple[str, dict[str, list[ExtractedField]]]:
    document = _python_docx.Document(io.BytesIO(data))
    blocks: list[str] = []
    for p in document.paragraphs:
        if p.text and p.text.strip():
            blocks.append(p.text.strip())
    for table in document.tables:
        rows = [[_cell_str(c.text) for c in row.cells] for row in table.rows]
        if rows:
            blocks.append(_markdown_table(rows[0], rows[1:]))
    return "\n\n".join(blocks).strip(), {}


def _parse_docx_zipxml(data: bytes, filename: str) -> tuple[str, dict[str, list[ExtractedField]]]:
    try:
        zf = zipfile.ZipFile(io.BytesIO(data))
    except zipfile.BadZipFile as exc:
        raise ValueError(f"{filename} is not a valid .docx archive") from exc
    with zf:
        if "word/document.xml" not in zf.namelist():
            raise ValueError(f"{filename} is not a valid .docx (missing word/document.xml)")
        root = etree.fromstring(zf.read("word/document.xml"))
    blocks: list[str] = []
    body = root.find(_WML + "body")
    for child in list(body) if body is not None else []:
        if child.tag == _WML + "p":
            text = _docx_paragraph_text(child)
            if text:
                blocks.append(text)
        elif child.tag == _WML + "tbl":
            table = _docx_table_markdown(child)
            if table:
                blocks.append(table)
        elif child.tag == _WML + "sectPr":
            continue
    return "\n\n".join(blocks).strip(), {}


# --------------------------------------------------------------------------- #
# Email (.eml) and HTML
# --------------------------------------------------------------------------- #
def _message_part_text(msg: Message) -> str:
    for part in msg.walk():
        if part.get_content_type() == "text/plain":
            payload = part.get_payload(decode=True)
            if isinstance(payload, bytes):
                charset = part.get_content_charset() or "utf-8"
                try:
                    return payload.decode(charset)
                except (LookupError, UnicodeDecodeError):
                    return payload.decode("utf-8", errors="replace")
    for part in msg.walk():
        if part.get_content_type() == "text/html":
            payload = part.get_payload(decode=True)
            if isinstance(payload, bytes):
                return _html_to_text(payload.decode("utf-8", errors="replace"))
    return ""


def _html_to_text(html: str) -> str:
    if BeautifulSoup is not None:
        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "head"]):
            tag.decompose()
        return _collapse_whitespace(soup.get_text(" ", strip=True))
    return _collapse_whitespace(_HTML_TAG_RE.sub(" ", html))


def email_body_text(msg: Message) -> str:
    """Best-effort readable text from an email message (plain text, else HTML)."""
    return _collapse_whitespace(_message_part_text(msg))


def parse_eml_bytes(data: bytes, filename: str = "submission.eml") -> tuple[str, dict[str, list[ExtractedField]]]:
    """Extract headers + readable body from a raw .eml file."""
    msg = message_from_bytes(data)
    headers = [f"{key}: {value}" for key in ("Subject", "From", "To", "Date", "Reply-To") if (value := msg.get(key))]
    body = _collapse_whitespace(_message_part_text(msg))
    parts = [block for block in [*headers, body] if block]
    fields: dict[str, list[ExtractedField]] = {}
    if subject := msg.get("Subject"):
        fields["eml.subject"] = [_field("eml.subject", subject, filename)]
    return "\n".join(parts).strip(), fields


def parse_html_bytes(data: bytes, filename: str = "file.html") -> tuple[str, dict[str, list[ExtractedField]]]:
    """Strip an HTML page down to readable text."""
    return _html_to_text(data.decode("utf-8", errors="replace")), {}


# --------------------------------------------------------------------------- #
# Public dispatch
# --------------------------------------------------------------------------- #
STRUCTURED_PARSERS: dict[str, Any] = {
    ".xlsx": parse_workbook_bytes,
    ".xls": parse_workbook_bytes,
    ".xlsm": parse_workbook_bytes,
    ".xltx": parse_workbook_bytes,
    ".xltm": parse_workbook_bytes,
    ".csv": parse_csv_bytes,
    ".docx": parse_docx_bytes,
    ".eml": parse_eml_bytes,
    ".msg": parse_eml_bytes,
    ".html": parse_html_bytes,
    ".htm": parse_html_bytes,
    ".mhtml": parse_html_bytes,
}

_TARGETED_STRUCTURED_EXTS = frozenset({".xlsx", ".xls", ".xlsm", ".xltx", ".xltm", ".csv", ".docx", ".eml", ".msg", ".html", ".htm", ".mhtml"})


def parse_structured_document(data: bytes, filename: str) -> tuple[str, str, dict[str, list[ExtractedField]]] | None:
    """Parse a binary/structured document, returning ``(text, engine, fields)``.

    Returns ``None`` when the extension has no structured parser, so callers can
    fall back to OCR or a raw text decode.
    """
    suffix = Path(filename).suffix.lower()
    parser = STRUCTURED_PARSERS.get(suffix)
    if parser is None:
        return None
    text, fields = parser(data, filename)
    return text, f"structured:{parser.__name__}", fields


def is_structured_extension(filename: str) -> bool:
    return Path(filename).suffix.lower() in _TARGETED_STRUCTURED_EXTS
